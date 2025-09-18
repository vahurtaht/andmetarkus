# impordi andmed REST teenusest, moodusta graafik ja salvestada tulemus wordi
import requests
# pip install python-docx
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

# Assuming 'df' is your DataFrame with the results
# If not, replace 'df' with your actual DataFrame variable

url_rahvastik_ = 'https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL?format=json'
url_women = "https://api.worldbank.org/v2/country/EST/indicator/SP.POP.TOTL.FE.IN?format=json"

response = requests.get(url_rahvastik_)
data = response.json()

response_woman = requests.get(url_women)
data_woman = response_woman.json()

values = {'year': [], 'population': [], 'women_population': []}


for item in data[1]:
    values['year'].append(item['date'])
    values['population'].append(item['value'])

for women in data_woman[1]:
    values['women_population'].append(women['value'])

df = pd.DataFrame(values)

df = df.sort_values(by='year')  # sorteerin aasta järgi kasvavalt


# Create a new Word document
doc = Document()
doc.add_heading('Eesti rahvaarv aastatel 1960-2021', level=1)
p = doc.add_paragraph(
    'Graafikul on näha Eesti rahvaarvu ja naiste arvu aastatel 1960-2021.')


# Joonista graafik ja salvesta pildina
df['men_population'] = df['population'] - df['women_population']
df_plot = df.sort_values(by='year')
plt.figure(figsize=(8, 4))
df_plot.plot(x='year', y=['population',
             'women_population', 'men_population'], kind='line')
plt.title('Eesti rahvaarv aastatel 1960-2021')
plt.xlabel('Aasta')
plt.ylabel('Inimeste arv')
plt.tight_layout()
plt.savefig('rahvaarv_plot.png')
plt.close()

# Lisa graafik Wordi dokumendile
doc.add_picture('rahvaarv_plot.png')


# add page break
doc.add_page_break()

doc.add_heading('Lisa: Algandmed', level=2)
doc.add_paragraph(
    'Allikas: The World Bank (https://data.worldbank.org/indicator/SP.POP.TOTL?locations=EE)')
doc.add_paragraph(
    'Allikas: The World Bank (https://data.worldbank.org/indicator/SP.POP.TOTL.FE.IN?locations=EE)')

# Add a table to the document
table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = str(col_name)

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)


# Save the document
doc.save("rahvaarv_tulemus.docx")
